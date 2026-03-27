import os
import re
import sys
import json
import traceback
import pandas as pd
from typing import Optional
from dotenv import load_dotenv

# ========== CARREGAR .env ==========

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(BASE_DIR)

dotenv_paths = [
    os.path.join(PROJECT_ROOT, ".env"),
    os.path.join(BASE_DIR, ".env"),
]

for dotenv_path in dotenv_paths:
    if os.path.exists(dotenv_path):
        load_dotenv(dotenv_path=dotenv_path, override=False)
        print(f"Carregado .env de: {dotenv_path}")

# ======= Dependências para DOCX e PDF =======

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except:
    Document = None

try:
    import comtypes.client as comtypes_client
except:
    comtypes_client = None

# ========= Gemini SDK =========

try:
    from google import genai
except:
    genai = None


# ========== DIRETÓRIOS ==========

DOWNLOAD_DIR = os.path.join(BASE_DIR, "downloads")
LIMPOS_DIR = os.path.join(DOWNLOAD_DIR, "arquivos limpos")
RESUMOS_DIR = os.path.join(DOWNLOAD_DIR, "resumos")
REG_FINAL = os.path.join(DOWNLOAD_DIR, "registros_icr.xlsx")

os.makedirs(DOWNLOAD_DIR, exist_ok=True)
os.makedirs(LIMPOS_DIR, exist_ok=True)
os.makedirs(RESUMOS_DIR, exist_ok=True)


# ========== FUNÇÕES AUXILIARES ==========

def codigo_para_nome_arquivo(codigo: str) -> str:
    return re.sub(r"[^\w\-]", "_", str(codigo))


def sindicato_eh_motorista(sindicatos: str) -> bool:
    if not sindicatos:
        return False
    return "MOTORIST" in sindicatos.upper()


# ========== EXTRAIR TEXTO VIA WORD (DOC / DOCX → TXT) ==========

def extrair_texto_arquivo(caminho: str) -> Optional[str]:
    if not os.path.exists(caminho):
        print(f"Arquivo não encontrado: {caminho}")
        return None

    _, ext = os.path.splitext(caminho)
    ext = ext.lower()

    if comtypes_client is None:
        print("comtypes não instalado — não é possível ler DOC.")
        return None

    try:
        print("  Exportando texto via Word...")

        word = comtypes_client.CreateObject("Word.Application")
        word.Visible = False

        base, _ = os.path.splitext(caminho)
        txt_tmp = base + "_tmp.txt"

        try:
            doc = word.Documents.Open(caminho)
            doc.SaveAs(txt_tmp, FileFormat=2)  # TXT
            doc.Close()
        finally:
            word.Quit()

        # Leitura
        try:
            with open(txt_tmp, "r", encoding="utf-8", errors="ignore") as f:
                texto = f.read()
        except:
            with open(txt_tmp, "r", encoding="latin-1", errors="ignore") as f:
                texto = f.read()

        os.remove(txt_tmp)

        texto = texto.strip()
        if not texto:
            print("  Arquivo convertido ficou vazio.")
            return None

        print(f"  Texto extraído ({len(texto)} chars)")
        return texto

    except Exception as e:
        print(f"Erro extraindo texto: {e}")
        traceback.print_exc()
        return None


# ========== CRIAR CLIENTE GEMINI ==========

def criar_cliente_gemini():
    if genai is None:
        print("google-genai não instalado.")
        return None

    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("GEMINI_API_KEY não encontrada.")
        return None

    try:
        client = genai.Client(api_key=api_key)
        return client
    except Exception as e:
        print(f"Erro criando cliente Gemini: {e}")
        return None


# ========== RESUMIR TODAS AS CLÁUSULAS (ESTILO 2) ==========

def resumir_instrumento_todas_clausulas(client, texto: str) -> Optional[str]:
    """
    Retorna um texto pronto para PDF contendo TODAS as cláusulas resumidas,
    numeradas em lista simples.
    """

    if not texto:
        return None

    max_len = 15000
    if len(texto) > max_len:
        texto = texto[:max_len]

    prompt = (
        "Você é um assistente especializado em resumir instrumentos coletivos.\n"
        "A seguir está o texto completo de uma convenção/acordo coletivo.\n\n"
        "TAREFA:\n"
        "- Leia TODO o texto.\n"
        "- Identifique TODAS as cláusulas.\n"
        "- Gere um RESUMO COMPLETO contendo:\n"
        "    1. Partes envolvidas\n"
        "    2. Vigência\n"
        "    3. Todas as cláusulas numeradas em LISTA SIMPLES, assim:\n"
        "       1. Cláusula X – Título: resumo breve.\n"
        "       2. Cláusula Y – Título: resumo breve.\n"
        "       ...\n\n"
        "FORMATAÇÃO:\n"
        "- NÃO USE markdown.\n"
        "- Apenas texto puro.\n"
        "- Seja objetivo e claro.\n"
        "- Não invente cláusulas inexistentes.\n"
        "- Mantenha os títulos originais das cláusulas sempre que possível.\n\n"
        "Agora processe o texto completo abaixo:\n\n"
    )

    try:
        resp = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt + texto,
        )

        resumo = (resp.text or "").strip()
        if not resumo:
            return None

        return resumo

    except Exception as e:
        print(f"Erro ao chamar Gemini: {e}")
        traceback.print_exc()
        return None


# ========== GERAR PDF A PARTIR DO RESUMO ==========

def salvar_resumo_pdf(codigo: str, resumo_texto: str) -> Optional[str]:

    if Document is None or comtypes_client is None:
        print("Dependências ausentes: python-docx / comtypes.")
        return None

    doc = Document()

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Resumo do Instrumento {codigo}")
    r.bold = True
    r.font.size = Pt(20)

    doc.add_paragraph("")

    # Corpo do resumo
    for linha in resumo_texto.split("\n"):
        par = doc.add_paragraph()
        run = par.add_run(linha.strip())
        run.font.size = Pt(11)

    nome_base = codigo_para_nome_arquivo(codigo)
    tmp_docx = os.path.join(RESUMOS_DIR, f"{nome_base}_tmp.docx")
    pdf_final = os.path.join(RESUMOS_DIR, f"{nome_base}_resumo.pdf")

    try:
        doc.save(tmp_docx)

        word = comtypes_client.CreateObject("Word.Application")
        word.Visible = False
        try:
            d = word.Documents.Open(tmp_docx)
            d.SaveAs(pdf_final, FileFormat=17)
            d.Close()
        finally:
            word.Quit()

        os.remove(tmp_docx)
        return pdf_final

    except Exception as e:
        print(f"Erro ao gerar PDF: {e}")
        traceback.print_exc()
        return None


# ========== MAIN ==========

def main():
    print("\n=== Resumo com Gemini (todas as cláusulas) ===")

    if not os.path.exists(REG_FINAL):
        print(f"Arquivo não encontrado: {REG_FINAL}")
        sys.exit(1)

    df = pd.read_excel(REG_FINAL, dtype=str).fillna("")

    if "resumido" not in df.columns:
        df["resumido"] = "não"
    if "resumo_arquivo" not in df.columns:
        df["resumo_arquivo"] = ""

    pendentes = df[
        df["resumido"].str.lower().ne("sim")
        & df["sindicatos"].apply(sindicato_eh_motorista)
    ].copy()

    print(f"Registros pendentes: {len(pendentes)}")

    if len(pendentes) == 0:
        print("Nada a fazer.")
        return

    client = criar_cliente_gemini()
    if client is None:
        print("Erro ao criar cliente Gemini.")
        return

    for _, row in pendentes.iterrows():

        codigo = str(row["codigo"])
        nome_base = codigo_para_nome_arquivo(codigo)

        caminho_doc = os.path.join(LIMPOS_DIR, f"{nome_base}.doc")
        caminho_docx = os.path.join(LIMPOS_DIR, f"{nome_base}.docx")

        if not os.path.exists(caminho_doc):
            caminho_doc = caminho_docx if os.path.exists(caminho_docx) else None

        if not caminho_doc:
            print(f"{codigo}: Arquivo não encontrado.")
            continue

        print(f"\nProcessando {codigo}...")

        texto = extrair_texto_arquivo(caminho_doc)
        if not texto:
            print("  Texto vazio. Pulando.")
            continue

        resumo_texto = resumir_instrumento_todas_clausulas(client, texto)
        if not resumo_texto:
            print("  Erro no resumo. Pulando.")
            continue

        pdf = salvar_resumo_pdf(codigo, resumo_texto)
        if not pdf:
            print("  Erro ao gerar PDF. Pulando.")
            continue

        print(f"  PDF gerado: {pdf}")

        idx = df.index[df["codigo"] == codigo]
        if len(idx):
            i = idx[0]
            df.at[i, "resumido"] = "sim"
            df.at[i, "resumo_arquivo"] = pdf

    df.to_excel(REG_FINAL, index=False)
    print("\nProcesso concluído.")


if __name__ == "__main__":
    main()
