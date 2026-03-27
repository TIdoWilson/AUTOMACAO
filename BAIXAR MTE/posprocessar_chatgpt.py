import os
import re
import sys
import json
import traceback
from typing import Optional

import pandas as pd
from dotenv import load_dotenv

# ================== CARREGAR .env ==================

# pasta onde este script está (ex.: .../python/BAIXAR MTE)
BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# pasta raiz do projeto (um nível acima: .../python)
PROJECT_ROOT = os.path.dirname(BASE_DIR)

# tenta carregar .env na raiz do projeto e na pasta do script
dotenv_paths = [
    os.path.join(PROJECT_ROOT, ".env"),
    os.path.join(BASE_DIR, ".env"),
]

for dotenv_path in dotenv_paths:
    if os.path.exists(dotenv_path):
        load_dotenv(dotenv_path=dotenv_path, override=False)
        print(f"Carregado .env de: {dotenv_path}")

# ---- imports para leitura e criação de documentos ----
try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError:
    Document = None

# comtypes para automação do Word (abrir .doc, exportar .txt, gerar PDF)
try:
    import comtypes.client as comtypes_client
except ImportError:
    comtypes_client = None

# ---- imports para OpenAI ----
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None


# ================== CONFIGURAÇÕES GERAIS ==================

DOWNLOAD_DIR = os.path.join(BASE_DIR, "downloads")
LIMPOS_DIR = os.path.join(DOWNLOAD_DIR, "arquivos limpos")
RESUMOS_DIR = os.path.join(DOWNLOAD_DIR, "resumos")

REG_FINAL = os.path.join(DOWNLOAD_DIR, "registros_icr.xlsx")

os.makedirs(DOWNLOAD_DIR, exist_ok=True)
os.makedirs(LIMPOS_DIR, exist_ok=True)
os.makedirs(RESUMOS_DIR, exist_ok=True)


# ================== FUNÇÕES AUXILIARES ==================

def codigo_para_nome_arquivo(codigo: str) -> str:
    """Converte o código em um nome seguro de arquivo."""
    return re.sub(r"[^\w\-]", "_", str(codigo))


def sindicato_eh_motorista(sindicatos: str) -> bool:
    """Verifica se o texto de sindicatos se refere a motoristas."""
    if not sindicatos:
        return False
    s = sindicatos.upper()
    return "MOTORIST" in s


def extrair_texto_arquivo(caminho: str) -> Optional[str]:
    """
    Extrai texto de .doc ou .docx.

    Preferência:
      - Se comtypes_client estiver disponível, usa o próprio Word para
        exportar o conteúdo como .txt (pega texto de parágrafos, tabelas,
        caixas de texto, etc.).
      - Caso contrário, para .docx tenta usar python-docx como fallback.
    """
    if not os.path.exists(caminho):
        print(f"  Arquivo nao encontrado para extracao: {caminho}")
        return None

    _, ext = os.path.splitext(caminho)
    ext = ext.lower()

    # 1) Caminho mais robusto: usar o Word para gerar .txt
    if ext in [".doc", ".docx"] and comtypes_client is not None:
        try:
            print("  Extraindo texto via Word (exportar para .txt)...")

            # abre o Word via COM
            word = comtypes_client.CreateObject("Word.Application")
            word.Visible = False

            base, _ = os.path.splitext(caminho)
            caminho_txt_tmp = base + "_texto_tmp.txt"

            try:
                doc_wd = word.Documents.Open(caminho)
                # 2 = wdFormatText (arquivo texto .txt)
                doc_wd.SaveAs(caminho_txt_tmp, FileFormat=2)
                doc_wd.Close()
            finally:
                word.Quit()

            texto = ""
            try:
                # tenta UTF-8 primeiro, ignorando erros se houver caracteres estranhos
                with open(caminho_txt_tmp, "r", encoding="utf-8", errors="ignore") as f:
                    texto = f.read()
            except UnicodeDecodeError:
                # fallback para latin-1, se necessário
                with open(caminho_txt_tmp, "r", encoding="latin-1", errors="ignore") as f:
                    texto = f.read()
            finally:
                try:
                    if os.path.exists(caminho_txt_tmp):
                        os.remove(caminho_txt_tmp)
                except OSError:
                    pass

            texto = (texto or "").strip()
            if not texto:
                print(f"  Nenhum texto util encontrado em {os.path.basename(caminho)}")
                return None

            print(f"  Texto extraido de {os.path.basename(caminho)} ({len(texto)} caracteres)")
            return texto

        except Exception as e:
            print(f"  Erro extraindo texto via Word de {caminho}: {e}")
            traceback.print_exc()
            return None

    # 2) Fallback: se for .docx e nao tivermos comtypes, tenta python-docx
    if ext == ".docx" and Document is not None:
        try:
            print("  Extraindo texto de .docx com python-docx (fallback)...")
            doc = Document(caminho)
            blocos = []

            for p in doc.paragraphs:
                txt = (p.text or "").strip()
                if txt:
                    blocos.append(txt)

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        txt = (cell.text or "").strip()
                        if txt:
                            blocos.append(txt)

            texto = "\n".join(blocos).strip()
            if not texto:
                print(f"  Nenhum texto util encontrado em {os.path.basename(caminho)}")
                return None

            print(f"  Texto extraido de {os.path.basename(caminho)} ({len(texto)} caracteres)")
            return texto
        except Exception as e:
            print(f"  Erro extraindo texto de {caminho}: {e}")
            traceback.print_exc()
            return None

    # 3) Extensao nao suportada ou sem ferramentas
    print(f"  Extensao nao suportada para extracao ou dependencias ausentes: {ext}")
    return None


def criar_cliente_openai() -> Optional["OpenAI"]:
    """Cria o cliente OpenAI usando a OPENAI_API_KEY do ambiente."""
    if OpenAI is None:
        print("Biblioteca 'openai' nao esta instalada.")
        return None

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print("OPENAI_API_KEY nao encontrada (nem no .env, nem no ambiente).")
        return None

    try:
        client = OpenAI(api_key=api_key)
        return client
    except Exception as e:
        print(f"Erro ao criar cliente OpenAI: {e}")
        return None


def resumir_instrumento_json(client: "OpenAI", texto: str) -> Optional[dict]:
    """
    Envia o texto do instrumento coletivo para o modelo e retorna um dict
    com os campos: partes, vigencia, economicas, outras.
    """
    if not texto:
        return None

    # para evitar textos muito gigantes (principalmente se vier .txt enorme)
    max_len = 12000
    if len(texto) > max_len:
        texto = texto[:max_len]

    prompt_sistema = (
        "Voce e um assistente que resume acordos coletivos de trabalho no Brasil. "
        "Resuma o instrumento de forma objetiva, em portugues, separando em quatro campos: "
        "1) partes envolvidas; 2) vigencia; 3) clausulas economicas; 4) outras clausulas relevantes para motoristas. "
        "Responda ESTRITAMENTE no seguinte formato JSON, sem texto extra, sem comentarios, sem markdown:\n\n"
        "{\n"
        '  \"partes\": \"...\",\n'
        '  \"vigencia\": \"...\",\n'
        '  \"economicas\": \"...\",\n'
        '  \"outras\": \"...\"\n'
        "}\n\n"
        "Seja conciso, use poucas frases em cada campo, apenas sintetize o conteudo sem copiar clausulas inteiras."
    )

    try:
        resp = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": prompt_sistema},
                {"role": "user", "content": texto},
            ],
            max_tokens=800,
        )
        conteudo = resp.choices[0].message.content
        if not conteudo:
            return None

        conteudo = conteudo.strip()

        # tenta decodificar diretamente como JSON
        try:
            data = json.loads(conteudo)
        except json.JSONDecodeError:
            # se vier com algum texto extra, tenta pegar apenas o bloco {...}
            m = re.search(r"\{.*\}", conteudo, flags=re.DOTALL)
            if not m:
                print("  Nao foi possivel decodificar JSON do resumo.")
                return None
            data = json.loads(m.group(0))

        return {
            "partes": str(data.get("partes", "")).strip(),
            "vigencia": str(data.get("vigencia", "")).strip(),
            "economicas": str(data.get("economicas", "")).strip(),
            "outras": str(data.get("outras", "")).strip(),
        }

    except Exception as e:
        print(f"  Erro ao chamar OpenAI: {e}")
        traceback.print_exc()
        return None


def salvar_resumo_pdf(codigo: str, resumo_dict: dict) -> Optional[str]:
    """
    Gera um DOCX de resumo bem formatado e converte para PDF usando Word.
    """
    if Document is None:
        print("  python-docx nao instalado; nao consigo gerar documento de resumo.")
        return None
    if comtypes_client is None:
        print("  comtypes nao instalado; nao consigo converter DOCX para PDF.")
        return None

    doc = Document()

    # Título centralizado
    titulo_par = doc.add_paragraph()
    titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo_par.add_run(f"Resumo do Instrumento {codigo}")
    titulo_par.style = "Title"
    titulo_run.bold = True
    titulo_run.font.size = Pt(20)

    # Linha em branco
    doc.add_paragraph("")

    # Pequeno subtítulo com explicação
    subtitulo_par = doc.add_paragraph()
    subtitulo_par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitulo_run = subtitulo_par.add_run(
        "Resumo estruturado do acordo coletivo, com foco em cláusulas relevantes para motoristas."
    )
    subtitulo_run.italic = True
    subtitulo_run.font.size = Pt(11)

    doc.add_paragraph("")

    # Seções organizadas
    secoes = [
        ("1) Partes envolvidas", resumo_dict.get("partes", "")),
        ("2) Vigência", resumo_dict.get("vigencia", "")),
        ("3) Cláusulas econômicas", resumo_dict.get("economicas", "")),
        ("4) Outras cláusulas relevantes para motoristas", resumo_dict.get("outras", "")),
    ]

    for titulo, conteudo in secoes:
        # Cabeçalho da seção
        p_titulo = doc.add_paragraph()
        p_titulo.style = "Heading 1"
        run_titulo = p_titulo.add_run(titulo)
        run_titulo.bold = True

        # Parágrafo com o texto
        p_texto = doc.add_paragraph()
        p_texto.style = "Normal"
        conteudo_limpo = " ".join(str(conteudo or "").split())
        if conteudo_limpo:
            run_texto = p_texto.add_run(conteudo_limpo)
        else:
            run_texto = p_texto.add_run("(sem informações relevantes)")
        run_texto.font.size = Pt(11)

        doc.add_paragraph("")

    nome_base = codigo_para_nome_arquivo(codigo)
    caminho_docx = os.path.join(RESUMOS_DIR, f"{nome_base}_resumo_tmp.docx")
    caminho_pdf = os.path.join(RESUMOS_DIR, f"{nome_base}_resumo.pdf")

    # Salva DOCX
    try:
        doc.save(caminho_docx)
    except Exception as e:
        print(f"  Erro ao salvar DOCX de resumo em {caminho_docx}: {e}")
        return None

    # Converte para PDF usando Word
    try:
        word = comtypes_client.CreateObject("Word.Application")
        word.Visible = False
        try:
            doc_wd = word.Documents.Open(caminho_docx)
            # 17 = wdFormatPDF
            doc_wd.SaveAs(caminho_pdf, FileFormat=17)
            doc_wd.Close()
        finally:
            word.Quit()
    except Exception as e:
        print(f"  Erro ao converter DOCX em PDF ({caminho_docx} -> {caminho_pdf}): {e}")
        traceback.print_exc()
        return None
    finally:
        # remove DOCX temporário
        try:
            if os.path.exists(caminho_docx):
                os.remove(caminho_docx)
        except OSError:
            pass

    if os.path.exists(caminho_pdf):
        return caminho_pdf
    else:
        return None


# ================== MAIN ==================

def main():
    print("=== Pos-processamento ChatGPT (motoristas -> PDF) ===")

    if not os.path.exists(REG_FINAL):
        print(f"Arquivo de registros nao encontrado: {REG_FINAL}")
        sys.exit(1)

    df = pd.read_excel(REG_FINAL, dtype=str)
    df.fillna("", inplace=True)

    # garante colunas de controle
    if "resumido" not in df.columns:
        df["resumido"] = "não"
    if "resumo_arquivo" not in df.columns:
        df["resumo_arquivo"] = ""

    # filtra apenas motoristas ainda nao resumidos
    mask_pendente = df["resumido"].str.lower().ne("sim") & df["sindicatos"].apply(sindicato_eh_motorista)
    df_pendentes = df[mask_pendente].copy()

    total_pendentes = len(df_pendentes)
    print(f"Registros pendentes para motoristas (nao resumidos): {total_pendentes}")

    if total_pendentes == 0:
        print("Nada para resumir. Encerrando.")
        df.to_excel(REG_FINAL, index=False)
        return

    client = criar_cliente_openai()
    if client is None:
        print("Nao foi possivel criar cliente OpenAI. Encerrando.")
        return

    for _, row in df_pendentes.iterrows():
        codigo = str(row["codigo"])
        sindicatos = str(row["sindicatos"])
        print(f"\nProcessando codigo {codigo} (sindicatos: {sindicatos[:80]}...)")

        nome_base = codigo_para_nome_arquivo(codigo)

        # tenta primeiro .doc, depois .docx
        caminho_doc = os.path.join(LIMPOS_DIR, f"{nome_base}.doc")
        if not os.path.exists(caminho_doc):
            caminho_docx = os.path.join(LIMPOS_DIR, f"{nome_base}.docx")
            if os.path.exists(caminho_docx):
                caminho_doc = caminho_docx
            else:
                print(f"  Nenhum arquivo .doc ou .docx encontrado para {codigo} (esperado: {nome_base}.doc/.docx)")
                continue

        print(f"  Arquivo encontrado: {os.path.basename(caminho_doc)}")

        texto = extrair_texto_arquivo(caminho_doc)
        if not texto:
            print("  Nao foi possivel extrair texto do arquivo. Pulando.")
            continue

        resumo_dict = resumir_instrumento_json(client, texto)
        if not resumo_dict:
            print("  Nao foi possivel obter resumo estruturado do ChatGPT. Pulando.")
            continue

        caminho_pdf = salvar_resumo_pdf(codigo, resumo_dict)
        if not caminho_pdf:
            print("  Falha ao salvar PDF de resumo. Pulando.")
            continue

        print(f"  Resumo salvo em: {caminho_pdf}")

        # Atualiza o dataframe original
        idx_real = df.index[df["codigo"] == codigo]
        if len(idx_real) > 0:
            i0 = idx_real[0]
            df.at[i0, "resumido"] = "sim"
            df.at[i0, "resumo_arquivo"] = caminho_pdf
        else:
            # fallback, caso tenha duplicidade estranha
            try:
                df.loc[df["codigo"] == codigo, "resumido"] = "sim"
                df.loc[df["codigo"] == codigo, "resumo_arquivo"] = caminho_pdf
            except Exception:
                pass

    df.to_excel(REG_FINAL, index=False)
    print("\nAtualizacao de registros concluida.")
    print(f"Arquivo atualizado: {REG_FINAL}")


if __name__ == "__main__":
    main()
